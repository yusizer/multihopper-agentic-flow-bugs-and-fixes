"""Convert findings/DRAFT.md -> SUBMISSION.html (for Google Docs import) + SUBMISSION.docx."""
import os, re, html
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
DRAFT = os.path.join(BASE, "findings", "DRAFT.md")
HTML_OUT = os.path.join(BASE, "findings", "SUBMISSION.html")
DOCX_OUT = os.path.join(BASE, "findings", "SUBMISSION.docx")

md = open(DRAFT, encoding="utf-8").read()

# ---------- HTML ----------
body = markdown.markdown(md, extensions=["tables", "fenced_code", "sane_lists", "nl2br"])
html_doc = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<title>Break It Before Users Do — MultiHopper Agentic Flow Bugs & Fixes</title>
<style>
body{{font-family:'Segoe UI',Arial,sans-serif;max-width:920px;margin:40px auto;padding:0 20px;line-height:1.55;color:#1a1a1a;}}
h1{{color:#9945FF;font-size:28px;}}h2{{color:#14F195;border-bottom:2px solid #14F195;padding-bottom:4px;margin-top:28px;}}
h3{{color:#0a7a5a;margin-top:22px;}}h4{{color:#444;}}
table{{border-collapse:collapse;width:100%;margin:12px 0;}}th,td{{border:1px solid #ccc;padding:6px 10px;text-align:left;vertical-align:top;}}
th{{background:#f4f0ff;}}tr:nth-child(even){{background:#fafafa;}}
code{{background:#f4f4f4;padding:2px 5px;border-radius:3px;font-family:Consolas,'Courier New',monospace;font-size:0.92em;}}
pre{{background:#0b0e16;color:#cbd5e1;padding:14px 16px;border-radius:8px;overflow:auto;}}pre code{{background:none;color:#cbd5e1;padding:0;}}
strong{{color:#111;}}hr{{border:none;border-top:1px solid #ddd;margin:24px 0;}}
</style></head><body>{body}</body></html>"""
open(HTML_OUT, "w", encoding="utf-8").write(html_doc)
print("HTML:", HTML_OUT, len(html_doc), "bytes")

# ---------- DOCX (basic markdown parse) ----------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Segoe UI"
style.font.size = Pt(10.5)

H_COLOR = {1: RGBColor(0x99, 0x45, 0xFF), 2: RGBColor(0x14, 0xF1, 0x95), 3: RGBColor(0x0a, 0x7a, 0x5a), 4: RGBColor(0x44, 0x44, 0x44)}

def add_runs(par, text):
    # **bold** and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for p in parts:
        if p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2]); r.bold = True
        elif p.startswith("`") and p.endswith("`"):
            r = par.add_run(p[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            par.add_run(p)

lines = md.split("\n")
i = 0
in_code = False
while i < len(lines):
    line = lines[i]
    if line.strip().startswith("```"):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Pt(12)
        r = par.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9)
        i += 1
        continue
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        lvl = len(m.group(1))
        h = doc.add_heading(level=lvl)
        run = h.add_run(m.group(2))
        run.font.color.rgb = H_COLOR.get(lvl, RGBColor(0, 0, 0))
        i += 1
        continue
    if line.strip().startswith("|") and "|" in line.strip()[1:]:
        # collect table rows
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(lines[i].strip()); i += 1
        # parse
        parsed = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
        # drop separator row (---)
        parsed = [r for r in parsed if not all(set(c) <= set("-: ") for c in r)]
        if parsed:
            t = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(parsed):
                for ci, cell in enumerate(row):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = ""
                        par = t.rows[ri].cells[ci].paragraphs[0]
                        add_runs(par, cell)
        continue
    if line.strip() == "" or line.strip() == "---":
        if line.strip() == "---":
            doc.add_paragraph("—").alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    # bullet/number
    mb = re.match(r"^[-*]\s+(.*)", line)
    mn = re.match(r"^\d+\.\s+(.*)", line)
    if mb:
        par = doc.add_paragraph(style="List Bullet"); add_runs(par, mb.group(1))
    elif mn:
        par = doc.add_paragraph(style="List Number"); add_runs(par, mn.group(1))
    else:
        par = doc.add_paragraph(); add_runs(par, line)
    i += 1

doc.save(DOCX_OUT)
print("DOCX:", DOCX_OUT)
